from aiogram import Router
from aiogram.filters import CommandStart
from aiogram.types import Message, CallbackQuery
from aiogram.fsm.state import StatesGroup, State
from aiogram import Router, F
from aiogram.types import Message
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.context import FSMContext
from aiogram.filters import StateFilter
from aiogram import BaseMiddleware
from aiogram.types import InputMediaVideo, InputMediaDocument
from aiogram.types import TelegramObject, FSInputFile
from aiogram import types
from aiogram.exceptions import TelegramBadRequest, TelegramRetryAfter, TelegramNetworkError
import re

import datetime
import asyncio
import traceback
from typing import Dict, Any, Callable, Awaitable, Optional
import time
import math

from docx import Document

import keyboard
import database
import config

db = database.MySQL()
CHAT_OWNER_CACHE: Dict[int, Dict[str, Any]] = {}


async def resolve_chat_owner(bot, chat_id: int, chat_name: str = "") -> Optional[int]:
    now_ts = time.time()
    cached = CHAT_OWNER_CACHE.get(chat_id)

    if cached is not None and cached.get("expires_at", 0) > now_ts:
        return cached.get("owner_id")

    try:
        chat_admins = await bot.get_chat_administrators(chat_id, request_timeout=7)
    except TelegramRetryAfter as e:
        # На flood-control ставим временную заглушку в кеш,
        # чтобы не долбить Telegram на каждом новом сообщении.
        CHAT_OWNER_CACHE[chat_id] = {
            "owner_id": cached.get("owner_id") if cached is not None else None,
            "expires_at": now_ts + max(int(e.retry_after), 5),
        }
        return CHAT_OWNER_CACHE[chat_id]["owner_id"]
    except TelegramNetworkError:
        # Временные сетевые ошибки Telegram (timeout/connect reset).
        # Ставим короткий cooldown, чтобы не блокировать обработку апдейтов.
        CHAT_OWNER_CACHE[chat_id] = {
            "owner_id": cached.get("owner_id") if cached is not None else None,
            "expires_at": now_ts + 30,
        }
        return CHAT_OWNER_CACHE[chat_id]["owner_id"]
    except TelegramBadRequest:
        # Если Telegram дал flood-control или чат не подходит для запроса,
        # просто пропускаем этот апдейт без исключения.
        CHAT_OWNER_CACHE[chat_id] = {
            "owner_id": None,
            "expires_at": now_ts + 60,
        }
        return None

    owner_id = None
    for admin in chat_admins:
        mentor_data = db.get_mentor_by_id(admin.user.id)
        if mentor_data is not None and chat_name.__contains__(mentor_data["mentor_name"]):
            owner_id = admin.user.id
            break

    CHAT_OWNER_CACHE[chat_id] = {
        "owner_id": owner_id,
        "expires_at": now_ts + 600,  # 10 минут
    }

    return owner_id


async def add_user_to_spreadsheet(user_id: int, email: str, flow: str, bot):
    """Добавляет пользователя в Google Таблицу если его там нет"""
    import gspread_asyncio
    from google.oauth2.service_account import Credentials
    import random
    
    try:
        # Проверяем, не добавлен ли уже пользователь
        is_in_added = db.is_email_in_added_api_users(email)
        if is_in_added:
            return  # Уже добавлен
        
        db.add_email_to_added_api_users(email)
        db.add_to_link_access(str(user_id), email.lower().strip(), flow)
        
        # Выбираем случайный chat для коммуникации (поддержка)
        support_chats = db.get_support_chats()
        if support_chats:
            random_support = random.choice(support_chats)
            communication_chat_id = random_support['support_chat_id']
        else:
            communication_chat_id = ""
        
        # Добавляем в Google Таблицу
        creds = Credentials.from_service_account_file("credentials.json")
        scoped = creds.with_scopes([
            'https://www.googleapis.com/auth/spreadsheets',
            'https://www.googleapis.com/auth/drive'
        ])
        agcm = gspread_asyncio.AioGspreadServiceAccount(from_credentials=scoped)
        agc = await agcm.authorize()
        ss_2 = await agc.open_by_url(config.SPREADSHEET_URL_USERS)
        table = await ss_2.get_worksheet_by_id(0)
        await table.append_row([email.lower().strip(), -1002572458943, flow, communication_chat_id, -1003545567896], value_input_option="USER_ENTERED")
    except Exception as e:
        # Логируем ошибку только в консоль
        print(f"Ошибка при добавлении в таблицу: {email} - {e}")


class QuestionChecker:
    def __init__(self):
        # Паттерны для вопросов
        self.patterns = {
            'question_mark': r'\?$',  # Вопрос в конце
            'question_words': [
                r'^(как|что|где|когда|почему|зачем|кто|чей|сколько|чему|чем|кому|кого)',
                r'^(можно ли|стоит ли|нужно ли|надо ли|есть ли|знаете ли|поможете ли)',
                r'^(помогите|подскажите|посоветуйте|объясните|расскажите)',
                r'^(а\s+)?(не\s+)?(знает|поможет|подскажет|сделает|сможет)\s+ли',
                r'^(какой|какая|какое|какие|чей|чья|чьё|чьи)',
            ],
            'question_phrases': [
                r'(не\s+)?подскажите,?\s+пожалуйста',
                r'(не\s+)?поможете,?\s+пожалуйста',
                r'как\s+быть',
                r'что\s+делать',
                r'что\s+мне\s+делать',
                r'как\s+мне\s+быть',
            ]
        }
    
    def is_question(self, text: str) -> bool:
        """Основной метод проверки"""
        if not text or not isinstance(text, str):
            return False
        
        text = self._clean_text(text)
        
        # Проверка на вопросительный знак
        if re.search(self.patterns['question_mark'], text):
            return True
        
        # Проверка на вопросительные слова/фразы
        for pattern in self.patterns['question_words']:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # Проверка на вопросительные фразы в любом месте текста
        for pattern in self.patterns['question_phrases']:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _clean_text(self, text: str) -> str:
        """Очистка текста"""
        # Убираем лишние пробелы
        text = text.strip()
        # Заменяем множественные пробелы на один
        text = re.sub(r'\s+', ' ', text)
        return text
    
    def get_question_type(self, text: str) -> str:
        """Определяет тип вопроса"""
        if not self.is_question(text):
            return "не вопрос"
        
        text_lower = self._clean_text(text).lower()
        
        if '?' in text:
            if re.search(r'\?$', text):
                return "прямой вопрос (с '?')"
            return "вопрос с '?'"
        
        if re.search(r'^(помогите|подскажите)', text_lower):
            return "просьба о помощи"
        
        if re.search(r'^(как|почему)', text_lower):
            return "вопрос о способе/причине"
        
        if re.search(r'^(что|кто)', text_lower):
            return "вопрос об объекте"
        
        return "вопрос"
    

class SubMiddleware(BaseMiddleware):
    async def __call__(
        self,
        handler: Callable[[TelegramObject, Dict[str, Any]], Awaitable[Any]],
        event: TelegramObject,
        data: Dict[str, Any]
    ) -> Any:
        user_data = db.get_user(event.from_user.id)
        
        if len(user_data) == 0:
            return await event.answer("Чтобы продолжить, отправьте почту сообщением в чат 👇")

        # Не блокируем callback-обработчики флагом готовности,
        # чтобы не попадать в спам/флуд-контроль Telegram.
        return await handler(event, data)
    

class SecondSubMiddleware(BaseMiddleware):
    async def __call__(
        self,
        handler: Callable[[TelegramObject, Dict[str, Any]], Awaitable[Any]],
        event: TelegramObject,
        data: Dict[str, Any]
    ) -> Any:
        try:
            print(f"[MSG] uid={event.from_user.id} chat_id={event.chat.id} type={event.chat.type} text={event.text}")
        except Exception:
            pass

        user_data = db.get_user(event.from_user.id)
        
        if len(user_data) == 0 and event.text != '/skip_state' and event.chat.id > 0 and event.text != '/start' and 'state' not in data:
            return await event.answer("Отправьте вашу почту в ответ на приветственное сообщение 👇")
        
        
        # Не держим сообщение в ожидании готовности и не отправляем
        # системные ответы, чтобы избежать Flood control на /start.
        return await handler(event, data)
    

start_router = Router()
start_router.callback_query.middleware(SubMiddleware())
start_router.message.middleware(SecondSubMiddleware())
quesion_checker = QuestionChecker()

class GetAccess(StatesGroup):
    email = State()


class SendHomeWork(StatesGroup):
    homework = State()


async def edit_message(message: Message, text: str, reply_markup=None, parse_mode=None) -> None:
    try:
        msg = await message.edit_text(text=text, reply_markup=reply_markup, parse_mode=parse_mode)
    except:
        try:
            await message.delete()
        except:
            pass

        msg = await message.answer(text=text, reply_markup=reply_markup, parse_mode=parse_mode)

    return msg


async def send_media_group(chat_id: int, bot):
    media = [
        InputMediaVideo(
            media=FSInputFile("files/IMG_0174.MOV")
        ),
        InputMediaDocument(
            media=FSInputFile("files/Инструкция по чат боту.pdf")
        )
    ]
    await bot.send_media_group(chat_id, media)


@start_router.message(CommandStart())
async def command_start_handler(message: Message, state: FSMContext) -> None:
    print(f"[START] uid={message.from_user.id} chat_id={message.chat.id} chat_type={message.chat.type} text={message.text}")
    await state.clear()

    support_chat_ids = [int(support_chat["support_chat_id"]) for support_chat in db.get_support_chats() if support_chat.get("support_chat_id")]
    is_tracker_chat = db.is_tracker(message.chat.id) is not None
    is_support_chat = message.chat.id in support_chat_ids
    is_tracker_user = str(message.from_user.id) in [str(i) for i in config.MANUAL_TRACKER_USER_IDS]

    is_staff = int(message.from_user.id) in [int(i) for i in config.DASHBOARD_ALLOWED_USER_IDS]

    user_data = db.get_user(message.from_user.id)

    if len(user_data) == 0:
        # In groups/supergroups do not start email onboarding flow.
        # This prevents trackers/support from being stuck in GetAccess.email state.
        if message.chat.type != 'private':
            print(f"[START] skip non-private for new user uid={message.from_user.id}")
            return

        start_args = None if len(message.text.split()) == 1 else "".join(message.text.split()[1:])
        link_access_data = None if start_args is None else db.get_link_access_by_user_id(start_args)

        if link_access_data is None:
            await state.set_state(GetAccess.email)
            
            await message.answer('''Привет! 👋

Рады видеть тебя в боте обучения «Заработок на Reels»

Чтобы открыть тебе доступ, напиши, пожалуйста, свою электронную почту, которая была указана при покупке обучения. Мы проверим тебя в списке учеников и сразу подключим к материалам 🚀

Напиши свою почту 👇''', reply_markup=keyboard.main_keyboard(include_dashboards=is_staff))
            print(f"[START] ask email uid={message.from_user.id}")
            
            return
        elif len(link_access_data) != 0:
            try:
                await message.bot.send_message(config.LOG_CHAT_ID, f'Новый пользователь в боте (Через быструю ссылку): {message.from_user.full_name} @{message.from_user.username} (ID: {message.from_user.id})\nПочта: {link_access_data[0]["email"]}')
            except:
                pass

            db.add_user(message.from_user.id, link_access_data[0]['email'])
            # Добавляем пользователя в Google Таблицу
            try:
                await add_user_to_spreadsheet(message.from_user.id, link_access_data[0]['email'], db.get_flow_by_email(link_access_data[0]['email']), message.bot)
            except Exception as e:
                print(f"Ошибка при добавлении в Google Таблицу: {e}")
        else:
            await state.set_state(GetAccess.email)
            
            await message.answer('''Привет! 👋

Рады видеть тебя в боте обучения «Заработок на Reels»

Чтобы открыть тебе доступ, напиши, пожалуйста, свою электронную почту, которая была указана при покупке обучения. Мы проверим тебя в списке учеников и сразу подключим к материалам 🚀

Напиши свою почту 👇''', reply_markup=keyboard.main_keyboard(include_dashboards=is_staff))
            print(f"[START] ask email (bad quick link) uid={message.from_user.id}")
            
            await message.bot.send_message(config.LOG_CHAT_ID, f'Не удалось найти информацию по быстрой ссылке: {message.from_user.full_name} @{message.from_user.username} (ID: {message.from_user.id})\nАргументы: {start_args}')
            return
        
    # Временный safe-mode: отключаем проверку непрочитанных,
    # чтобы /start отвечал мгновенно без риска подвисаний на БД.
    has_tracker_unread = False

    welcome_text = """Рады видеть тебя в обучении «Заработок на Reels»📱

Здесь твои домашние задания, связь с личным трекером и поддержкой. Всё в одном месте — удобно и без лишних поисков

Выбирай нужный раздел ниже и нажимай на кнопку👇"""

    try:
        await message.answer_photo(
            photo=FSInputFile("files/start.jpg"),
            caption=welcome_text,
            reply_markup=keyboard.main_keyboard(include_dashboards=is_staff, has_tracker_unread=has_tracker_unread),
        )
        print(f"[START] welcome(photo) sent uid={message.from_user.id}")
    except TelegramRetryAfter as e:
        await asyncio.sleep(max(int(e.retry_after), 1))
        try:
            await message.answer(
                welcome_text,
                reply_markup=keyboard.main_keyboard(include_dashboards=is_staff, has_tracker_unread=has_tracker_unread),
            )
            print(f"[START] welcome(text after retry) sent uid={message.from_user.id}")
        except Exception:
            print(f"[START] failed after retry uid={message.from_user.id}\n{traceback.format_exc()}")
            pass
    except (TelegramNetworkError, FileNotFoundError):
        try:
            await message.answer(
                welcome_text,
                reply_markup=keyboard.main_keyboard(include_dashboards=is_staff, has_tracker_unread=has_tracker_unread),
            )
            print(f"[START] welcome(text fallback) sent uid={message.from_user.id}")
        except Exception:
            print(f"[START] fallback failed uid={message.from_user.id}\n{traceback.format_exc()}")
            pass


@start_router.callback_query(F.data == 'main')
async def command_start_handler(call: CallbackQuery, state: FSMContext) -> None:
    await state.clear()

    support_chat_ids = [int(support_chat["support_chat_id"]) for support_chat in db.get_support_chats() if support_chat.get("support_chat_id")]
    is_tracker_chat = db.is_tracker(call.message.chat.id) is not None
    is_support_chat = call.message.chat.id in support_chat_ids
    is_tracker_user = str(call.from_user.id) in [str(i) for i in config.MANUAL_TRACKER_USER_IDS]

    is_staff = int(call.from_user.id) in [int(i) for i in config.DASHBOARD_ALLOWED_USER_IDS]

    try:
        await call.message.delete()
    except:
        pass

    # Временный safe-mode: отключаем проверку непрочитанных,
    # чтобы возврат в меню работал стабильно.
    has_tracker_unread = False

    try:
        await call.message.answer_photo(photo=FSInputFile("files/start.jpg"), caption="""Рады видеть тебя в обучении «Заработок на Reels»📱

Здесь твои домашние задания, связь с личным трекером и поддержкой. Всё в одном месте — удобно и без лишних поисков

Выбирай нужный раздел ниже и нажимай на кнопку👇""", reply_markup=keyboard.main_keyboard(include_dashboards=is_staff, has_tracker_unread=has_tracker_unread))
    except TelegramRetryAfter:
        await call.message.answer("""Рады видеть тебя в обучении «Заработок на Reels»📱

Здесь твои домашние задания, связь с личным трекером и поддержкой. Всё в одном месте — удобно и без лишних поисков

Выбирай нужный раздел ниже и нажимай на кнопку👇""", reply_markup=keyboard.main_keyboard(include_dashboards=is_staff, has_tracker_unread=has_tracker_unread))


@start_router.callback_query(F.data == 'study_menu')
async def command_start_handler(call: CallbackQuery, state: FSMContext) -> None:
    await state.clear()
    await edit_message(call.message, 'Выберите действие из меню ниже:', reply_markup=keyboard.study_keyboard())


@start_router.callback_query(F.data == 'get_instruction')
async def command_start_handler(call: CallbackQuery) -> None:
    try:
        await call.message.delete()
    except:
        pass

    await call.message.answer_document(FSInputFile("files/Инструкция по чат боту.pdf"), reply_markup=keyboard.back_to_main_keyboard())


@start_router.callback_query(F.data == 'get_homeworks')
async def command_start_handler(call: CallbackQuery) -> None:
    await edit_message(call.message, 'Выбери тип домашнего задания', reply_markup=keyboard.get_homeworks_keyboard())


@start_router.callback_query(F.data == 'my_claps')
async def command_start_handler(call: CallbackQuery) -> None:
    user_data = db.get_user(call.from_user.id)
    users_flow = db.get_flow_by_email(user_data[0]['email'])

    if float(users_flow) < 14.3:
        await call.answer('Только для потоков 14.3 и выше')
        return
    
    done_homework_ids = db.get_done_homework_ids(call.from_user.id)
    required_homework = db.get_required_homework_ids(users_flow)
    count_required_homework = 0

    ignore_lessons = []

    for i in done_homework_ids:
        if i in required_homework and i not in ignore_lessons:
            count_required_homework += 1

            if "analog" in required_homework[i]:
                ignore_lessons += required_homework[i]["analog"]

    count_claps = 15
    claps_path = "files/"

    if float(users_flow) >= 15.1 and float(users_flow) < 15.6:
        count_claps = 12
        claps_path = "files/15_1/"
    elif float(users_flow) >= 15.6:
        count_claps = 11
        claps_path = "files/15_6/"

    if count_required_homework == 0:
        await edit_message(call.message, f"""Пока у тебя нет ни одной хлопушки «Стоп, снято!», но это поправимо! 🎬

Напомню правила: за каждое принятое <u>обязательное задание</u> ты получаешь одну хлопушку.

Собери все {count_claps} хлопушек, и я открою тебе доступ к подарку — Бонусному уроку «Тренды в контенте в 2026 году»! 🔥

Жду твои работы! 👏🏻""", parse_mode="HTML", reply_markup=keyboard.support_keyboard())
        return
    
    try:
        await call.message.delete()
    except:
        pass
    
    clap_string = 'хлопушек'

    if count_required_homework == 1:
        clap_string = 'хлопушка'
    elif count_required_homework >= 2 and count_required_homework <= 4:
        clap_string = 'хлопушки'

    await call.bot.send_photo(call.from_user.id, photo=FSInputFile(f'{claps_path}{count_required_homework}.jpg'), caption=f"""Крутой прогресс! У тебя уже <b>{count_required_homework}</b> {clap_string} «Стоп, снято!» 🎬

Продолжай в том же духе! Всего {count_claps} хлопушек — и бонусный урок о трендах 2026 года твой! 🔥""", parse_mode="HTML", reply_markup=keyboard.support_keyboard())


@start_router.callback_query(F.data.startswith('get_homeworks:'))
async def command_start_handler(call: CallbackQuery) -> None:
    homework_type = call.data.split(':')[1]
    print(homework_type)

    user_data = db.get_user(call.from_user.id)
    users_flow = db.get_flow_by_email(user_data[0]['email'])

    if len(call.data.split(':')) > 2:
        message_id = int(call.data.split(':')[2])

        try:
            await call.bot.delete_message(chat_id=call.from_user.id, message_id=message_id)
        except:
            pass

    homework_data = db.get_all_user_homeworks(call.from_user.id)
    homework_list = []
    done_homework = False

    if homework_type == 'checking':
        text = 'Список домашних заданий на проверке'

        for homework in homework_data:
            if homework["status"] == 'На проверке' or homework["status"] == '⏳':
                module_data = db.get_module(str(homework["module_id"]), users_flow)

                if module_data is None:
                    continue

                module = module_data["name"]
                lesson_data = db.get_lesson(str(homework["lesson_id"]), users_flow)

                if lesson_data is None:
                    continue

                homework_name = lesson_data["name"]
                homework_list.append({'name': f'{module} - {homework_name}', 'id': homework["lesson_id"]})

    elif homework_type == 'rework':
        text = 'Список домашних заданий на доработке'

        for homework in homework_data:
            if homework["status"] == '❌':
                module_data = db.get_module(str(homework["module_id"]), users_flow)

                if module_data is None:
                    continue

                module = module_data["name"]
                lesson_data = db.get_lesson(str(homework["lesson_id"]), users_flow)

                if lesson_data is None:
                    continue

                homework_name = lesson_data["name"]
                homework_list.append({'name': f'{module} - {homework_name}', 'id': homework["lesson_id"]})

    elif homework_type == 'sent':
        text = 'Список выполненных домашних заданий'
        done_homework = True

        for homework in homework_data:
            if homework["status"] == '✅':
                module_data = db.get_module(str(homework["module_id"]), users_flow)

                if module_data is None:
                    continue

                module = module_data["name"]
                lesson_data = db.get_lesson(str(homework["lesson_id"]), users_flow)

                if lesson_data is None:
                    continue

                homework_name = lesson_data["name"]
                homework_list.append({'name': f'{module} - {homework_name}', 'id': homework["lesson_id"]})

    elif homework_type == 'obligatory':
        text = 'Список обязательных домашних заданий:'
        done_lessons, rework_lessons, sent_lessons, check_lessons = [], [], [], []

        for homework in homework_data:
            if homework["status"] == '✅':
                done_lessons.append(homework["lesson_id"])

            if homework["status"] == '❌':
                rework_lessons.append(homework["lesson_id"])

            if homework["status"] == '⏳':
                check_lessons.append(homework["lesson_id"])

            if homework["status"] == '📤':
                sent_lessons.append(homework["lesson_id"])

        homework_list_2 = []
        ignore_required_homework_ids = []
        required_homework = db.get_required_homework_ids(users_flow)

        for i in required_homework:
            lesson_data_2 = db.get_lesson(str(i), users_flow)

            if lesson_data_2 is None:
                continue

            if i in ignore_required_homework_ids:
                continue

            homework_data_2 = []
            homework_data_2.append({'name': lesson_data_2['name'], 'lesson_id': str(i)})

            if 'analog' in required_homework[i]:
                for j in required_homework[i]['analog']:
                    if j == i:
                        continue

                    lesson_data_2 = db.get_lesson(str(j), users_flow)

                    if lesson_data_2 is None:
                        continue

                    homework_data_2.append({'name': lesson_data_2['name'], 'lesson_id': str(j)})

                ignore_required_homework_ids += required_homework[i]['analog']

            homework_list_2.append(homework_data_2)

        await edit_message(call.message, text, reply_markup=keyboard.lessons_keyboard_2(homework_list_2, done_lessons, rework_lessons, check_lessons, sent_lessons), parse_mode='HTML')
        return

    if len(homework_list) == 0:
        text = 'Список домашних заданий пуст'

    await edit_message(call.message, text, reply_markup=keyboard.get_homeworks_list_keyboard(homework_list, done_homework, homework_type))


@start_router.callback_query(F.data == 'delete_message')
async def command_start_handler(call: CallbackQuery) -> None:
    await call.message.delete()


@start_router.callback_query(F.data.startswith('homework_is_done:'))
async def homework_is_done(call: CallbackQuery) -> None:
    lesson_id = call.data.split(':')[1]
    homework_data = db.get_homework_by_lesson_id(call.from_user.id, lesson_id)

    try:
        await call.message.delete()
    except:
        pass
    
    last_solution_text = await get_last_solution_text(call.from_user.id, lesson_id, call.from_user.username)

    if not last_solution_text:
        try:
            msg_1 = await call.bot.copy_message(chat_id=call.from_user.id, from_chat_id=homework_data[0]["chat_id"], message_id=homework_data[0]["message_id_2"])
        except Exception as e:
            print(f"ДЗ пользователя не найдено: {e} {call.from_user.id} {lesson_id}")
            msg_1 = await call.message.answer('Текст ДЗ не найден')
    else:
        msg_1 = await call.message.answer_document(FSInputFile(last_solution_text))

    if len(call.data.split(':')) > 2:
        reply_markup = keyboard.back_from_lesson_keyboard(homework_data[0]["module_id"], msg_1.message_id, False, f"get_homeworks:{call.data.split(':')[2]}:{msg_1.message_id}")
    else:
        reply_markup = keyboard.back_from_lesson_keyboard(homework_data[0]["module_id"], msg_1.message_id)

    # TODO рассмотреть подробнее
    try:
        await call.message.answer(homework_data[0]["comment"], reply_markup=reply_markup)
    except:
        document = Document()
        document.add_paragraph(homework_data[0]["comment"])
        file_name = f'cache/Ответ на ДЗ {call.from_user.id if call.from_user.username is None else "@" + call.from_user.username}.docx'
        document.save(file_name)

        await call.message.answer_document(FSInputFile(file_name), reply_markup=reply_markup)


@start_router.message(GetAccess.email)
async def command_start_handler(message: Message, state: FSMContext) -> None:
    is_staff = int(message.from_user.id) in [int(i) for i in config.DASHBOARD_ALLOWED_USER_IDS]

    # Prevent email onboarding flow from hijacking messages in group chats
    if message.chat.type != 'private':
        await state.clear()
        return

    if not message.text:
        await message.answer("Пожалуйста, введите email текстовым сообщением", reply_markup=keyboard.main_keyboard(include_dashboards=is_staff))
        return
    
    is_access = db.is_email_in_users_access(message.text.lower())

    if is_access:
        db.add_user(message.from_user.id, message.text.lower())
        await state.clear()
        
        users_flow = db.get_flow_by_email(message.text.lower())
        
        # Добавляем пользователя в Google Таблицу
        try:
            await add_user_to_spreadsheet(message.from_user.id, message.text.lower(), users_flow, message.bot)
        except Exception as e:
            print(f"Ошибка при добавлении в Google Таблицу: {e}")

        if float(users_flow) >= 14.3:
            await message.answer_document(FSInputFile('files/Инструкция по чат боту.pdf'))

        await message.answer('''Проверка твоего аккаунта прошла успешно, ты можешь приступить к выполнению домашних заданий.
Увидимся в рекомендациях!''', reply_markup=keyboard.main_keyboard())
        
        try:
            await message.bot.send_message(config.LOG_CHAT_ID, f'Новый пользователь в боте: {message.from_user.full_name} @{message.from_user.username} (ID: {message.from_user.id})\nПочта: {message.text.lower()}')
        except:
            pass
    else:
        await message.answer('''Не вижу тебя в списке учеников.

Что делать? 
1. Проверь правильно ли указана электронная почта и попробуй еще раз
2. Если уверен(а), что почта указана корректно, то обязательно напиши нам в поддержку http://t.me/stepbybit''', reply_markup=keyboard.main_keyboard(include_dashboards=is_staff))
        
        try:
            await message.bot.send_message(config.LOG_CHAT_ID, f'Ученик не может авторизоваться: {message.from_user.full_name} @{message.from_user.username} (ID: {message.from_user.id})\nПочта: {message.text.lower()}')
        except:
            pass


@start_router.callback_query(F.data == 'get_onboarding')
async def command_start_handler(call: CallbackQuery) -> None:
    user_data = db.get_user(call.from_user.id)
    email_key = user_data[0]["email"].lower().strip()

    if email_key not in config.USERS_ADDITIONAL_INFO or config.USERS_ADDITIONAL_INFO[email_key]["tariff"] is None or len(config.USERS_ADDITIONAL_INFO[email_key]["tariff"]) == 0:
        await call.answer("У вас не указан тариф, обратитесь в поддержку", show_alert=True)
    elif config.USERS_ADDITIONAL_INFO[email_key]["tariff"] == "Про повышение до VIP": # TODO указать номер потока
        await edit_message(call.message, """Привет-привет! 🙌
Я Иляна из команды заботы Арины Алекс 💛
Поздравляю с покупкой менторства “Заработок на Reels”, тариф ВИП, поток 15.7 - впереди вас ждет незабываемое приключение! 🎬

Вот ссылка на платформу: https://school.stepbybit.ru/

Чтобы зайти, нажмите «Восстановить пароль», укажите свою почту, и на неё придёт письмо со входом.

Начните с Вводного модуля - там короткие уроки, домашнее задание, стартовая анкета (по ней вас распределят в мини-группу) и договор оферты, чтобы открылись следующие уроки.

Ваши бонусы при покупке можно найти: Тренинги - Обучение по заработку на Рилс - Бонусы и библиотека материалов.

<b>В первый день обучения:</b> 
1. Вам придёт письмо на почту с ссылкой на мини-группу с ментором.

2. Ваш персональный трекер напишет в WhatsApp или Телеграмм - она будет помогать вам с заданиями и подсказывать, если что-то непонятно 🙂

Чтобы быть в курсе новостей и обновлений, обязательно подпишитесь на каналы в Телеграмме

https://t.me/addlist/kWP0mHTSa3IyMWJi

Рада, что вы с нами! Всё получится - предобучение стартует 16 февраля, основная программа - 23 февраля 🚀""", reply_markup=keyboard.back_to_main_keyboard(), parse_mode="HTML")
    elif config.USERS_ADDITIONAL_INFO[email_key]["tariff"] == "База повышение до Про": # TODO указать номер потока
        await edit_message(call.message, """Привет-привет! 🙌
Я Иляна из команды заботы Арины Алекс 💛
Поздравляю с покупкой менторства “Заработок на Reels”, тариф ПРО, поток 15.7 - впереди вас ждет незабываемое приключение! 🎬

Вот ссылка на платформу: https://school.stepbybit.ru/

Чтобы зайти, нажмите «Восстановить пароль», укажите свою почту, и на неё придёт письмо со входом.

Начните с Вводного модуля — там короткие уроки, домашнее задание, стартовая анкета (по ней вас распределят в мини-группу) и договор оферты, чтобы открылись следующие уроки.

Ваши бонусы при покупке можно найти: Тренинги - Обучение по заработку на Рилс - Бонусы и библиотека материалов.

<b>В первый день обучения:</b>
1. Вам придёт письмо на почту с ссылкой на мини-группу с ментором.

2. Ваш персональный трекер напишет в WhatsApp или Телеграмм - она будет помогать вам с заданиями и подсказывать, если что-то непонятно 🙂

Чтобы быть в курсе новостей и обновлений, обязательно подпишитесь на каналы в Телеграмм:

https://t.me/addlist/de2kQMGg21piOThi

Рада, что вы с нами! Всё получится - предобучение стартует 16 февраля, основная программа - 23 февраля 🚀""", reply_markup=keyboard.back_to_main_keyboard(), parse_mode="HTML")
    else:
        await call.answer("У вас не указан тариф, обратитесь в поддержку", show_alert=True)

@start_router.callback_query(F.data == 'get_modules')
async def command_start_handler(call: CallbackQuery) -> None:
    user_data = db.get_user(call.from_user.id)
    users_flow = db.get_flow_by_email(user_data[0]['email'])

    modules_list = db.get_modules(users_flow)

    await edit_message(call.message, 'Выбери модуль, чтобы выполнить задания:', reply_markup=keyboard.modules_keyboard(modules_list))
  

@start_router.callback_query(F.data.startswith('get_module:'))
async def command_start_handler(call: CallbackQuery, state: FSMContext) -> None:
    await state.clear()

    module_id = call.data.split(':')[1]
    user_data = db.get_user(call.from_user.id)
    users_flow = db.get_flow_by_email(user_data[0]['email'])
    module_access_data = db.get_module_access_2(users_flow, module_id)

    if len(module_access_data) != 0 and module_access_data[0]['time'] > time.time():
        date_string = datetime.datetime.fromtimestamp(module_access_data[0]['time']).strftime('%d.%m')
        await call.answer(f'Доступ к данному модулю будет доступен {date_string}', show_alert=True)
        return

    if len(call.data.split(':')) > 2:
        delete_message_id = call.data.split(':')[2]

        try:
            await call.message.bot.delete_message(call.from_user.id, delete_message_id)
        except:
            pass

    lessons_list = db.get_lessons(module_id, users_flow)
    module_data = db.get_module(module_id, users_flow)
    
    if module_data is None:
        error_message = f"⚠️ Ошибка при загрузке мини-аппа\n\nПользователь: {call.from_user.full_name} (@{call.from_user.username}, ID: {call.from_user.id})\nМодуль ID: {module_id}\nПоток: {users_flow}\n\nОшибка: module_data returned None\n\nTraceback:\n{traceback.format_exc()}"
        try:
            for admin_id in config.ADMINS_LIST:
                await call.message.bot.send_message(admin_id, error_message)
        except:
            pass
        await call.answer("Произошла ошибка при загрузке. Пожалуйста, сообщите об этом в поддержку.", show_alert=True)
        return
    
    module_name = module_data['name']
    module_desc = ("\n" + module_data['description']) if module_data['description'] != '' else ''

    rework_lessons = db.get_rework_lessons_ids(call.from_user.id, module_id)
    check_lessons = db.get_check_lessons_ids(call.from_user.id, module_id)
    done_lessons = db.get_done_lessons_ids(call.from_user.id, module_id)
    sent_lessons = db.get_sent_lessons_ids(call.from_user.id, module_id)

    await edit_message(call.message, f'{module_name}{module_desc}', reply_markup=keyboard.lessons_keyboard(lessons_list, done_lessons, rework_lessons, check_lessons, sent_lessons))


@start_router.callback_query(F.data.startswith('get_lesson:'))
async def command_start_handler(call: CallbackQuery, state: FSMContext) -> None:
    lesson_id = call.data.split(':')[1]

    homework_data = db.get_homework_by_lesson_id(call.from_user.id, lesson_id)

    user_data = db.get_user(call.from_user.id)
    users_flow = db.get_flow_by_email(user_data[0]['email'])

    lesson_data = db.get_lesson(lesson_id, users_flow)

    module_access_data = db.get_module_access_2(users_flow, lesson_data['module_id'])

    if len(module_access_data) != 0 and module_access_data[0]['time'] > time.time():
        date_string = datetime.datetime.fromtimestamp(module_access_data[0]['time']).strftime('%d.%m')
        await call.answer(f'Доступ к данному уроку будет доступен {date_string}', show_alert=True)
        return

    if len(homework_data) != 0 and homework_data[0]['status'] == '⏳':
        await call.answer('Вы отправили ДЗ на проверку', show_alert=True)
        return
    elif len(homework_data) != 0 and homework_data[0]['status'] == 'На проверке':
        await call.answer('ДЗ на проверке у трекера', show_alert=True)
        return
    elif len(homework_data) != 0 and homework_data[0]['status'] == '✅':
        await state.set_state(SendHomeWork.homework)
        await state.update_data(lesson_id=lesson_id)
        return await homework_is_done(call)

    if len(call.data.split(':')) > 2:
        reply_markup = keyboard.back_from_lesson_keyboard(module_id=lesson_data['module_id'], return_to_module=False, callback_data=f"get_homeworks:{call.data.split(':')[2]}", last_solution=True if len(homework_data) != 0 else False)
        await state.update_data(callback_data=f"get_homeworks:{call.data.split(':')[2]}")
    else:
        reply_markup = keyboard.back_from_lesson_keyboard(lesson_data['module_id'], last_solution=True if len(homework_data) != 0 else False)

    lesson_text = f'{lesson_data["name"]}\n\n{lesson_data["task_text"]}'

    msg = await edit_message(call.message, lesson_text, reply_markup=reply_markup)
    await state.set_state(SendHomeWork.homework)

    await state.update_data(lesson_id=lesson_id, message_id=msg.message_id, module_id=lesson_data['module_id'], lesson_text=lesson_text)


async def get_last_solution_text(tg_id, lesson_id, username): # Возвращает длинное_решение: bool или длинное_решение: str
    homework_text_data = db.get_homework_text_data(tg_id, lesson_id)

    if len(homework_text_data) == 0:
        return False

    last_solution_time = homework_text_data[0]['time']
    text_data = []
    
    for i in range(0, len(homework_text_data)):
        if last_solution_time - homework_text_data[i]['time'] > 10:
            continue

        text_data.append(homework_text_data[i]['text'])
    
    if len(text_data) == 1 or homework_text_data[0]['text'] is None:
        return False

    text_data.reverse()

    solution_text = '\n'.join(text_data)

    document = Document()
    document.add_paragraph(solution_text)
    file_name = f'cache/ДЗ от {tg_id if username is None else "@" + username}.docx'
    document.save(file_name)

    return file_name


@start_router.callback_query(StateFilter(SendHomeWork.homework), F.data == 'last_solution')
async def command_start_handler(call: CallbackQuery, state: FSMContext) -> None:
    try:
        await call.message.delete()
    except:
        pass

    state_data = await state.get_data()
    
    homework_data = db.get_homework_by_lesson_id(call.from_user.id, state_data['lesson_id'])
    last_solution_text = await get_last_solution_text(call.from_user.id, state_data['lesson_id'], call.from_user.username)

    if not last_solution_text:
        try:
            msg_1 = await call.bot.copy_message(chat_id=call.from_user.id, from_chat_id=homework_data[0]["chat_id"], message_id=homework_data[0]["message_id_2"])
        except Exception as e:
            print(f"ДЗ пользователя не найдено: {e} {call.from_user.id}")
            msg_1 = await call.message.answer('Текст ДЗ не найден')
    else:
        msg_1 = await call.message.answer_document(FSInputFile(last_solution_text))

    try:
        msg_2 = await call.message.answer(homework_data[0]["comment"], reply_markup=keyboard.last_solution_keyboard(msg_1.message_id))
    except:
        document = Document()
        document.add_paragraph(homework_data[0]["comment"])
        file_name = f'cache/Ответ на ДЗ {call.from_user.id if call.from_user.username is None else "@" + call.from_user.username}.docx'
        document.save(file_name)

        msg_2 = await call.message.answer_document(FSInputFile(file_name), reply_markup=keyboard.last_solution_keyboard(msg_1.message_id))

    await state.update_data(message_id=msg_2.message_id)


@start_router.callback_query(F.data.startswith('check_solution:'))
async def command_start_handler(call: CallbackQuery, state: FSMContext) -> None:
    try:
        await call.message.delete()
    except:
        pass

    lesson_id = call.data.split(':')[1]

    user_data = db.get_user(call.from_user.id)
    users_flow = db.get_flow_by_email(user_data[0]['email'])
    
    homework_data = db.get_homework_by_lesson_id(call.from_user.id, lesson_id)

    last_solution_text = await get_last_solution_text(call.from_user.id, lesson_id, call.from_user.username)

    if not last_solution_text:
        try:
            msg_1 = await call.bot.copy_message(chat_id=call.from_user.id, from_chat_id=homework_data[0]["chat_id"], message_id=homework_data[0]["message_id_2"])
        except Exception as e:
            print(f"ДЗ пользователя не найдено: {e} {call.from_user.id} {lesson_id}")
            msg_1 = await call.message.answer('Текст ДЗ не найден')
    else:
        msg_1 = await call.message.answer_document(FSInputFile(last_solution_text))
    
    try:
        msg_2 = await call.message.answer(homework_data[0]["comment"], reply_markup=keyboard.last_solution_keyboard(msg_1.message_id))
    except:
        document = Document()
        document.add_paragraph(homework_data[0]["comment"])
        file_name = f'cache/Ответ на ДЗ {call.from_user.id if call.from_user.username is None else "@" + call.from_user.username}.docx'
        document.save(file_name)

        msg_2 = await call.message.answer_document(FSInputFile(file_name), reply_markup=keyboard.last_solution_keyboard(msg_1.message_id))

    await state.set_state(SendHomeWork.homework)

    lesson_data = db.get_lesson(lesson_id, users_flow)
    
    lesson_text = f'{lesson_data["name"]}\n\n{lesson_data["task_text"]}'

    await state.update_data(lesson_id=lesson_id, message_id=msg_2.message_id, module_id=lesson_data['module_id'], lesson_text=lesson_text)


@start_router.callback_query(StateFilter(SendHomeWork.homework), F.data.startswith('back_from_last_solution:'))
async def command_start_handler(call: CallbackQuery, state: FSMContext) -> None:
    state_data = await state.get_data()
    delete_message_id = call.data.split(':')[1]

    try:
        await call.bot.delete_message(call.from_user.id, delete_message_id)
    except:
        pass

    if 'callback_data' in state_data:
        reply_markup = keyboard.back_from_lesson_keyboard(module_id=state_data['module_id'], return_to_module=False, callback_data=state_data['callback_data'], last_solution=True)
    else:
        reply_markup = keyboard.back_from_lesson_keyboard(state_data['module_id'], last_solution=True)

    await edit_message(call.message, state_data['lesson_text'], reply_markup=reply_markup)


@start_router.callback_query(F.data == 'get_strategy')
async def command_start_handler(call: CallbackQuery) -> None:
    user_data = db.get_user(call.from_user.id)
    users_flow = db.get_flow_by_email(user_data[0]['email'])

    if float(users_flow) < 15:
        await call.answer('Только для потоков 15 и выше')
        return
    
    await edit_message(call.message, 'Выберите стратегию:', reply_markup=keyboard.get_strategy_keyboard())


@start_router.callback_query(F.data == 'get_promotion_strategy')
async def command_start_handler(call: CallbackQuery) -> None:
    homework_text_data = db.get_homework_text_data(call.from_user.id, 16)

    if len(homework_text_data) == 0:
        await call.answer('Вы не сдали «Моя стратегия продвижения»')
        return

    last_solution_time = homework_text_data[0]['time']
    text_data = []
    
    for i in range(0, len(homework_text_data)):
        if last_solution_time - homework_text_data[i]['time'] > 10:
            continue

        text_data.append(homework_text_data[i]['text'])
    
    try:
        await call.message.delete()
    except:
        pass
    
    if len(text_data) == 1 or homework_text_data[0]['text'] is None:
        homework_data = db.get_homework_by_lesson_id(call.from_user.id, 16)
        await call.bot.copy_message(chat_id=call.from_user.id, from_chat_id=homework_data[0]["chat_id"], message_id=homework_data[0]["message_id_2"], reply_markup=keyboard.back_from_promotion_strategy_keyboard())
        return

    text_data.reverse()

    solution_text = '\n'.join(text_data)

    document = Document()
    document.add_paragraph(solution_text)
    file_name = f'cache/Моя стратегия продвижения {call.from_user.id if call.from_user.username is None else "@" + call.from_user.username}.docx'
    document.save(file_name)

    await call.message.answer_document(FSInputFile(file_name), reply_markup=keyboard.back_from_promotion_strategy_keyboard())


@start_router.callback_query(F.data == 'get_automatic_strategy')
async def command_start_handler(call: CallbackQuery) -> None:
    await edit_message(call.message, '''Автоматическая стратегия - это ваш индивидуальный файл, в котором автоматически собраны ваши выполненные домашние задания из соответсвующих разделов стратегии🔥

Для удобства вы можете воспользоваться данным файлом для формирования итоговой стратегии при выполнении домашнего задания «Моя стратегия продвижения»''', reply_markup=keyboard.get_automatic_strategy_keyboard())


@start_router.callback_query(F.data == 'automatic_strategy_file')
async def command_start_handler(call: CallbackQuery) -> None:
    try:
        await call.message.delete()
    except:
        pass
    
    file_name = f'cache/Автоматическая стратегия {call.from_user.id if call.from_user.username is None else "@" + call.from_user.username}.docx'

    document = Document()
    p = document.add_paragraph()

    for key, value in config.AUTOMATIC_STRATEGY_LESSONS.items():
        p.add_run(key).bold = True
        homework_text_data = db.get_homework_text_data(call.from_user.id, value)
        homework_data = db.get_homework_by_lesson_id(call.from_user.id, value)

        if len(homework_text_data) == 0 or homework_data[0]['status'] != '✅':
            p.add_run('\n\nЗадание не сдано\n\n')
            continue

        last_solution_time = homework_text_data[0]['time']
        text_data = []
        
        for i in range(0, len(homework_text_data)):
            if last_solution_time - homework_text_data[i]['time'] > 10:
                continue

            text_data.append(homework_text_data[i]['text'])
        
        if len(text_data) == 1 and homework_text_data[0]['text'] is not None:
            p.add_run(f'\n\n{homework_text_data[0]["text"]}\n\n')
            continue
        elif homework_text_data[0]['text'] is None:
            p.add_run('\n\nЗадание отправлено в медиа\n\n')
            continue

        text_data.reverse()
        solution_text = '\n'.join(text_data)
        p.add_run(f'\n\n{solution_text}\n\n')        

    document.save(file_name)

    await call.message.answer_document(FSInputFile(file_name), reply_markup=keyboard.back_from_automatic_strategy_keyboard())


@start_router.message(SendHomeWork.homework)
async def command_start_handler(message: Message, state: FSMContext) -> None:
    state_data = await state.get_data()
    
    # Check if lesson_id exists in state data
    if "lesson_id" not in state_data:
        await message.answer("Пожалуйста, выберите урок из меню для сдачи домашнего задания.")
        return
    
    user_data = db.get_user(message.from_user.id)
    users_flow = db.get_flow_by_email(user_data[0]['email'])

    homework_data = db.get_homework_by_lesson_id(message.from_user.id, state_data["lesson_id"])

    if len(homework_data) != 0 and homework_data[0]['status'] == '✅':
        try:
            await message.delete()
        except:
            pass
        return

    try:
        await message.bot.edit_message_reply_markup(chat_id=message.chat.id, message_id=state_data['message_id'], reply_markup=None)
    except:
        pass

    SEND_CHAT_ID = db.get_chat_id(user_data[0]["email"].lower())
    homework_name = (db.get_lesson(str(state_data["lesson_id"]), users_flow))["name"]

    is_do_homework = db.get_homework_by_lesson_id(message.from_user.id, state_data["lesson_id"])

    update_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    msg_1 = await message.bot.send_message(SEND_CHAT_ID, f'⬇️ {message.from_user.full_name} @{message.from_user.username} ({user_data[0]["email"].lower()} Поток: {users_flow}) сделал ДЗ! Урок: {homework_name} (Техническая информация: {message.from_user.id}_{state_data["lesson_id"]}) ⬇️')
    msg_2 = await message.bot.copy_message(SEND_CHAT_ID, message.chat.id, message.message_id)

    if len(is_do_homework) != 0:
        db.edit_homework(
            homework_id=is_do_homework[0]["homework_id"],
            status="⏳",
            comment="",
            update_time=update_time,
            message_link=f"https://t.me/c/{-(SEND_CHAT_ID+1000000000000)}/{msg_1.message_id}",
            message_id_1=msg_1.message_id,
            message_id_2=msg_2.message_id,
            check_time="",
            tg_id=message.from_user.id,
            send_message_id=message.message_id
        )
    else:
        db.add_homework(
            f"{message.from_user.full_name} @{message.from_user.username} {message.from_user.id}",
            state_data["lesson_id"],
            "⏳",
            "",
            update_time,
            f"https://t.me/c/{-(SEND_CHAT_ID+1000000000000)}/{msg_1.message_id}",
            "",
            msg_1.message_id,
            msg_2.message_id,
            message.from_user.id,
            state_data['module_id'],
            message.message_id,
            SEND_CHAT_ID
        )

    if 'callback_data' in state_data:
        reply_markup = keyboard.back_from_lesson_keyboard(state_data['module_id'], return_to_module=False, callback_data=state_data['callback_data'])
    else:
        reply_markup = keyboard.back_from_lesson_keyboard(state_data['module_id'])

    await message.answer('Твое задание отправлено трекеру на проверку✅\nОбратная связь поступит в этот бот в течении 2-3 дней', reply_markup=reply_markup)

    db.add_homework_text(message.from_user.id, state_data["lesson_id"], time.time(), message.text)

    await state.clear()


async def send_congratulation_message(message: Message, lesson_id: int, user_id):
    print(lesson_id, user_id)
    user_data = db.get_user(user_id)
    
    if len(user_data) == 0:
        return
    
    users_flow = db.get_flow_by_email(user_data[0]['email'])
    print(users_flow)

    if float(users_flow) < 14.3:
        return
    
    print(1)
    
    done_homework_ids = db.get_done_homework_ids(user_id)
    required_homework = db.get_required_homework_ids(users_flow)

    if lesson_id not in required_homework and lesson_id not in config.IMPORTANT_HOMEWORKS_IDS:
        return
    
    print(2)
    
    count_required_homework = 0
    ignore_lessons = []

    for i in done_homework_ids:
        if i in required_homework and i not in ignore_lessons:
            count_required_homework += 1

            if "analog" in required_homework[i]:
                ignore_lessons += required_homework[i]["analog"]

    lesson_required_data = required_homework.get(lesson_id, {})

    if "analog" in lesson_required_data:
        count_analog_done = 0

        for i in done_homework_ids:
            if i in lesson_required_data["analog"]:
                count_analog_done += 1

        if count_analog_done >= 2:
            return

    if count_required_homework >= 12 and float(users_flow) >= 15.1 and float(users_flow) < 15.6:
        lesson_data = db.get_lesson(str(lesson_id), users_flow)
        await message.bot.send_photo(user_id, photo=FSInputFile(f'files/15_1/{count_required_homework}.jpg'), caption="""Поздравляем!🎉 Тобой сданы все 12 обязательных заданий!🎬 Ты молодец😍

Забирай свой бонус - «Тренды в контенте в 2026 году»""", parse_mode="HTML")
        await asyncio.sleep(0.5)
        await message.bot.send_document(user_id, FSInputFile(f'files/Тренды в контенте 2026.pdf'))

    elif count_required_homework >= 11 and float(users_flow) >= 15.6:
        lesson_data = db.get_lesson(str(lesson_id), users_flow)
        await message.bot.send_photo(user_id, photo=FSInputFile(f'files/15_6/{count_required_homework}.jpg'), caption="""Поздравляем!🎉 Тобой сданы все 11 обязательных заданий!🎬 Ты молодец😍

Забирай свой бонус - «Тренды в контенте в 2026 году»""", parse_mode="HTML")
        await asyncio.sleep(0.5)
        await message.bot.send_document(user_id, FSInputFile(f'files/Тренды в контенте 2026.pdf'))

    elif count_required_homework >= 15 and float(users_flow) < 15.1:
        lesson_data = db.get_lesson(str(lesson_id), users_flow)
        await message.bot.send_photo(user_id, photo=FSInputFile(f'files/{count_required_homework}.jpg'), caption="""Поздравляем!🎉 Тобой сданы все 15 обязательных заданий!🎬 Ты молодец😍

Забирай свой бонус - «Тренды в контенте в 2026 году»""", parse_mode="HTML")
        await asyncio.sleep(0.5)
        await message.bot.send_document(user_id, FSInputFile(f'files/Тренды в контенте 2026.pdf'))
    
    elif float(users_flow) < 15.1:
        lesson_data = db.get_lesson(str(lesson_id), users_flow)
        await message.bot.send_photo(user_id, photo=FSInputFile(f'files/{count_required_homework}.jpg'), caption=config.CONGRATULATION_MESSAGE.replace("<lesson_name>", lesson_data['name']), parse_mode="HTML")

    elif float(users_flow) >= 15.1 and float(users_flow) < 15.6:
        lesson_data = db.get_lesson(str(lesson_id), users_flow)
        await message.bot.send_photo(user_id, photo=FSInputFile(f'files/15_1/{count_required_homework}.jpg'), caption=config.CONGRATULATION_MESSAGE_2.replace("<lesson_name>", lesson_data['name']), parse_mode="HTML")
    elif float(users_flow) >= 15.6:
        lesson_data = db.get_lesson(str(lesson_id), users_flow)
        await message.bot.send_photo(user_id, photo=FSInputFile(f'files/15_6/{count_required_homework}.jpg'), caption=config.CONGRATULATION_MESSAGE_3.replace("<lesson_name>", lesson_data['name']), parse_mode="HTML")


@start_router.message(F.text == '/my_chat_id')
async def my_chat_id(message: types.Message):
    await message.answer(
        f"Ваш user_id: {message.from_user.id}\n"
        f"Текущий chat_id: {message.chat.id}\n"
        f"Тип чата: {message.chat.type}"
    )


@start_router.message(F.text.startswith('/flapper'))
async def flapper(message: types.Message):
    if message.from_user.id not in config.ADMINS_LIST:
        return
    elif len(message.text.split()) != 3:
        return await message.answer('Неправильная команда. Отправьте: /flapper TG_ID ID_урока или /flapper @username ID_урока или /flapper email@example.com ID_урока')
    
    user_param = message.text.split()[1]
    lesson_id = message.text.split()[2]
    user_data = []
    
    # Если передан @username, пробуем получить ID через Telegram API
    if user_param.startswith('@'):
        try:
            chat = await message.bot.get_chat(user_param)
            user_id = str(chat.id)
        except Exception as e:
            return await message.answer(f'Пользователь {user_param} не найден. Убедитесь, что бот когда-либо взаимодействовал с этим пользователем.')
    elif '@' in user_param:
        # Поиск по email
        user_data = db.get_user_by_email(user_param.lower().strip())
        if len(user_data) == 0:
            return await message.answer('Пользователь не найден в базе данных')

        tg_id = user_data[0].get('tg_id')
        if tg_id is None or not str(tg_id).isdigit() or int(tg_id) == 0:
            return await message.answer('У пользователя в базе не указан Telegram ID. Попросите пользователя написать боту /start, затем повторите команду.')

        user_id = str(tg_id)
    else:
        user_id = user_param
    
    if len(user_data) == 0:
        user_data = db.get_user(user_id)

    if len(user_data) == 0:
        return await message.answer('Пользователь не найден в базе данных')

    try:
        users_flow = db.get_flow_by_email(user_data[0]['email'])
    except Exception:
        return await message.answer('Для пользователя не найден поток обучения (users_access).')

    lesson_data = db.get_lesson(lesson_id, users_flow)

    if lesson_data is None:
        return await message.answer('Урок не найден')

    await send_congratulation_message(message, int(lesson_id), user_id)
    await message.answer("Хлопушка успешно отправлена!")


@start_router.message(F.text.startswith('/fixuser'))
async def fixuser(message: types.Message):
    """Команда для обновления данных пользователя (tg_id, username)"""
    if message.from_user.id not in config.ADMINS_LIST:
        return
    
    parts = message.text.split()
    
    if len(parts) == 1:
        # /fixuser - обновить данные текущего пользователя
        user_id = str(message.from_user.id)
        username = message.from_user.username
        
        user_data = db.get_user(user_id)
        if len(user_data) == 0:
            return await message.answer('Вы не найдены в базе данных. Сначала авторизуйтесь в боте.')
        
        # Обновляем username
        if username:
            db.update_user_username(user_id, username)
            await message.answer(f'Ваш username @{username} обновлен в базе данных!')
        else:
            await message.answer('У вас нет username в Telegram, обновить невозможно.')
        return
    
    if len(parts) != 2:
        return await message.answer('Неправильная команда. Используйте:\n/fixuser - обновить свои данные\n/fixuser @username - обновить данные по username\n/fixuser email@example.com - обновить данные по email')
    
    user_param = parts[1]
    current_user_id = str(message.from_user.id)
    current_username = message.from_user.username
    
    # Пытаемся найти пользователя по @username или email
    if user_param.startswith('@'):
        # Ищем по username в базе
        user_data = db.get_user_by_username(user_param[1:])
        if user_data is None or len(user_data) == 0:
            return await message.answer(f'Пользователь с username {user_param} не найден в базе данных.')
    elif '@' in user_param:
        user_data = db.get_user_by_email(user_param.lower())
        if len(user_data) == 0:
            return await message.answer(f'Пользователь с email {user_param} не найден в базе данных.')
    else:
        return await message.answer('Неверный формат. Используйте @username или email.')
    
    # Обновляем tg_id и username
    target_tg_id = user_data[0]['tg_id'] if 'tg_id' in user_data[0] else None
    
    # Обновляем данные
    db.update_user_tg_id(user_data[0]['email'], current_user_id)
    if current_username:
        db.update_user_username(current_user_id, current_username)
    
    await message.answer(f'Данные обновлены!\nEmail: {user_data[0]["email"]}\nTG ID: {current_user_id}\nUsername: @{current_username}')


@start_router.message(F.text.startswith('/finduser'))
async def finduser(message: types.Message):
    """Команда для поиска пользователя по email или username и вывода информации"""
    if message.from_user.id not in config.ADMINS_LIST:
        return
    
    parts = message.text.split()
    
    if len(parts) != 2:
        return await message.answer('Используйте: /finduser email@example.com или /finduser @username')
    
    user_param = parts[1]
    user_data = None
    
    # Ищем по username или email
    if user_param.startswith('@'):
        user_data = db.get_user_by_username(user_param[1:])
    elif '@' in user_param:
        user_data = db.get_user_by_email(user_param.lower())
    else:
        return await message.answer('Используйте: /finduser email@example.com или /finduser @username')
    
    if user_data is None or len(user_data) == 0:
        return await message.answer(f'Пользователь {user_param} не найден в базе данных.')
    
    user = user_data[0]
    email = user.get('email', 'не указан')
    tg_id = user.get('tg_id', 'не указан')
    username = user.get('username', 'не указан')
    
    # Получаем информацию о потоке
    try:
        users_flow = db.get_flow_by_email(email)
    except:
        users_flow = 'не найден'
    
    # Проверяем, есть ли пользователь в USERS_ADDITIONAL_INFO
    tracker_chat_id = config.USERS_ADDITIONAL_INFO.get(email, {}).get('tracker_chat_id', 'не назначен')
    
    info_text = f"""Информация о пользователе:

Email: {email}
TG ID: {tg_id}
Username: @{username}
Поток: {users_flow}
Chat трекера: {tracker_chat_id}

Команды для обновления:
/fixtgid {email} - обновить TG ID текущего пользователя
/fixuser {email} - обновить все данные (tg_id + username)"""
    
    await message.answer(info_text)


@start_router.message(
    (F.text.startswith('/fixtracker') | F.text.startswith('/repairtracker'))
)
async def fix_tracker_user_by_email(message: types.Message):
    """
    Самоисправление проблемного ученика для трекеров.
    Команда пытается восстановить/синхронизировать tg_id по email
    (users <-> link_access), чтобы чат трекера снова открывался.
    """
    allowed_ids = set(int(i) for i in config.MANUAL_TRACKER_USER_IDS)
    allowed_ids.add(5201430878)

    tracker_ids = set()
    try:
        tracker_ids = set(int(i) for i in db.get_trackers_chats())
    except Exception:
        tracker_ids = set()

    if (
        int(message.from_user.id) not in allowed_ids
        and int(message.from_user.id) not in tracker_ids
        and int(message.from_user.id) not in [int(i) for i in config.ADMINS_LIST]
    ):
        return

    parts = message.text.split(maxsplit=1)
    if len(parts) != 2:
        return await message.answer(
            "Используйте: /fixtracker email@example.com\n"
            "или: /repairtracker email@example.com"
        )

    email = parts[1].strip().lower()
    if '@' not in email or ' ' in email:
        return await message.answer("Неверный формат email. Пример: /fixtracker ilona_76@bk.ru")

    users_rows = db.get_user_by_email(email)
    link_rows = db.get_link_access_by_email(email)

    valid_user_ids = []
    for row in users_rows:
        tg_id = row.get('tg_id')
        if tg_id is not None and str(tg_id).isdigit() and int(tg_id) > 0:
            valid_user_ids.append(int(tg_id))

    link_user_ids = []
    for row in link_rows:
        link_user_id = row.get('user_id')
        if link_user_id is not None and str(link_user_id).isdigit() and int(link_user_id) > 0:
            link_user_ids.append(int(link_user_id))

    chosen_tg_id = None

    # Приоритет: уже валидный tg_id в users, иначе берем самый свежий из link_access
    if len(valid_user_ids) != 0:
        chosen_tg_id = valid_user_ids[0]
    elif len(link_user_ids) != 0:
        chosen_tg_id = link_user_ids[-1]

    if chosen_tg_id is None:
        return await message.answer(
            "Не удалось починить автоматически: нет валидного Telegram ID в users/link_access.\n"
            "Попросите ученика написать боту /start и повторите команду."
        )

    try:
        if len(users_rows) == 0:
            db.add_user(chosen_tg_id, email)
            action_text = f"Создал запись users: tg_id={chosen_tg_id}"
        else:
            db.update_user_tg_id(email, chosen_tg_id)
            action_text = f"Синхронизировал tg_id в users: {chosen_tg_id}"

        open_link = f"https://rb.infinitydev.tw1.su/get_tracker_chat?user_id={chosen_tg_id}"
        await message.answer(
            "✅ Исправление выполнено\n"
            f"Email: {email}\n"
            f"{action_text}\n"
            f"users rows: {len(users_rows)}\n"
            f"link_access rows: {len(link_rows)}\n"
            f"Открыть чат: {open_link}"
        )
    except Exception as e:
        await message.answer(
            "⚠️ Ошибка при исправлении.\n"
            f"Email: {email}\n"
            f"Текст ошибки: {e}"
        )


@start_router.message_reaction()
async def message_reaction_handler(message_reaction: types.MessageReactionUpdated) -> Any:
    support_chats_list = db.get_support_chats()
    support_chat_ids = [int(support_chat["support_chat_id"]) for support_chat in support_chats_list]
    chat_ids_list = db.get_chat_ids()

    if message_reaction.chat.id in support_chat_ids or message_reaction.chat.id == config.PSYHOLOGIST_CHAT_ID or message_reaction.chat.id in chat_ids_list:
        return
    
    message_data = db.get_chat_message(message_reaction.chat.id, message_reaction.message_id)

    if message_data is None:
        return

    owner_id = await resolve_chat_owner(
        message_reaction.bot,
        message_reaction.chat.id,
        message_reaction.chat.full_name or message_reaction.chat.title or "",
    )

    if owner_id is None:
        return
    
    db.change_count_reactions(message_reaction.chat.id, message_reaction.message_id, message_data["count_reactions"] + 1 if len(message_reaction.new_reaction) > message_data["count_reactions"] else message_data["count_reactions"] - 1)


@start_router.message(StateFilter(None))
async def command_start_handler(message: Message) -> None:
    support_chats_list = db.get_support_chats()
    support_chat_ids = [int(support_chat["support_chat_id"]) for support_chat in support_chats_list]
    chat_ids_list = db.get_chat_ids()
    trackers_chats = db.get_trackers_chats()

    # Обработка системы метрик
    if message.chat.type in ('group', 'supergroup') and message.chat.id not in support_chat_ids and message.chat.id != config.PSYHOLOGIST_CHAT_ID and message.chat.id not in chat_ids_list and str(message.chat.id) not in trackers_chats:
        chat_type = None
        owner_id = await resolve_chat_owner(
            message.bot,
            message.chat.id,
            message.chat.full_name or message.chat.title or "",
        )
        if owner_id is not None:
            chat_type = 'mentor'
        
        if owner_id is not None:
            is_question = False if message.text is None or message.text == "" else quesion_checker.is_question(message.text)

            db.add_chat_message(owner_id, message.chat.id, message.message_id, 0, chat_type, message.from_user.id, 0, is_question if datetime.datetime.now().time().hour >= 10 and datetime.datetime.now().time().hour < 19 else False, int(time.time()), 0)

            if message.reply_to_message is not None:
                message_data = db.get_chat_message(message.reply_to_message.chat.id, message.reply_to_message.message_id)

                if message_data is None:
                    return

                db.change_reply_count(message.reply_to_message.chat.id, message.reply_to_message.message_id, message_data["reply_count"] + 1)

                # Если это ответ на вопрос TODO уточнить момент по рабочему времени
                if message_data["is_question"] and message_data["tg_id"] != owner_id and owner_id == message.from_user.id:
                    db.change_unix_time_answered(message.reply_to_message.chat.id, message.reply_to_message.message_id, int(time.time()))

        return
    
    # Обработка системы выгрузки у трекеров
    if str(message.chat.id) in trackers_chats and message.text is not None and message.text == '/list':
        await message.reply("Список учеников трекера:", reply_markup=keyboard.web_app_tracker_list_keyboard(message.chat.id))
        return
    
    # Обработка системы трекеров
    if str(message.chat.id) in trackers_chats and message.reply_to_message is not None and message.reply_to_message.text is not None and message.reply_to_message.text.__contains__('(Техническая информация:'):
        user_id = int(message.reply_to_message.text.split("Техническая информация: ")[-1].split(")")[0])
        
        try:
            await message.bot.copy_message(
                user_id, 
                message.chat.id, 
                message.message_id if (message.text is None or not message.text.startswith('/msg_id')) else int(message.text.split()[1]), 
                reply_markup=keyboard.tracker_keyboard_2()
            )

            await message.bot.set_message_reaction(message.chat.id, message.message_id, [{"type": "emoji", "emoji": "👍"}])

            file_id = None
            file_type = None

            if message.photo is not None:
                file_id = message.photo[-1].file_id
                file_type = 'photo'
            elif message.video is not None:
                file_id = message.video.file_id
                file_type = 'video'
            elif message.document:
                file_id = message.document.file_id
                file_type = 'document'
            elif message.audio:
                file_id = message.audio.file_id
                file_type = 'audio'
            elif message.voice:
                file_id = message.voice.file_id
                file_type = 'voice'
            elif message.video_note:
                file_id = message.video_note.file_id
                file_type = 'video_note'

            db.add_to_trackers_messages(message.from_user.id, message.chat.id, message.html_text, file_id, file_type, False, time.time(), f"https://t.me/c/{-(message.chat.id+1000000000000)}/{message.message_id}")

            # ✅ сервер рассылает событие (источник истины)
            message_payload = {
                "type": "message",
                "text": message.html_text,
                "sender_id": "0",
                "unix_time": int(time.time()),
            }

            # отправляем ВСЕМ подключённым (включая отправителя)
            for ws in config.ws_connections[str(user_id)]:
                await ws.send_json(message_payload)
        except Exception as e:
            print(e)
            pass

        return

    # Обработка сообщений из чата с психологом
    if message.chat.id == config.PSYHOLOGIST_CHAT_ID and message.reply_to_message is not None and message.reply_to_message.text is not None and message.reply_to_message.text.__contains__('(Техническая информация:'):
        user_id = int(message.reply_to_message.text.split("Техническая информация: ")[-1].split(")")[0])
        
        try:
            await message.bot.copy_message(
                user_id, 
                message.chat.id, 
                message.message_id if (message.text is None or not message.text.startswith('/msg_id')) else int(message.text.split()[1]), 
                reply_markup=keyboard.psychologist_keyboard_2()
            )

            await message.bot.set_message_reaction(message.chat.id, message.message_id, [{"type": "emoji", "emoji": "👍"}])

            file_id = None
            file_type = None

            if message.photo is not None:
                file_id = message.photo[-1].file_id
                file_type = 'photo'
            elif message.video is not None:
                file_id = message.video.file_id
                file_type = 'video'
            elif message.document:
                file_id = message.document.file_id
                file_type = 'document'
            elif message.audio:
                file_id = message.audio.file_id
                file_type = 'audio'
            elif message.voice:
                file_id = message.voice.file_id
                file_type = 'voice'
            elif message.video_note:
                file_id = message.video_note.file_id
                file_type = 'video_note'

            db.add_to_psychologist_messages(message.from_user.id, message.chat.id, message.html_text, file_id, file_type, False, time.time(), f"https://t.me/c/{-(message.chat.id+1000000000000)}/{message.message_id}")

            # ✅ сервер рассылает событие (источник истины)
            message_payload = {
                "type": "message",
                "text": message.html_text,
                "sender_id": "0",
                "unix_time": int(time.time()),
            }

            # отправляем ВСЕМ подключённым (включая отправителя)
            for ws in config.ws_connections_psychologist[str(user_id)]:
                await ws.send_json(message_payload)
        except Exception as e:
            print(e)
            pass

        return
    
    elif message.chat.id == config.PSYHOLOGIST_CHAT_ID and message.reply_to_message is not None:
        await message.reply(str(message.reply_to_message.message_id))
        return

    if message.chat.id in support_chat_ids and message.reply_to_message is not None and message.reply_to_message.text is not None and message.reply_to_message.text.__contains__('(Техническая информация:'):
        user_id = int(message.reply_to_message.text.split("Техническая информация: ")[-1].split(")")[0])
        
        try:
            await message.bot.copy_message(
                user_id, 
                message.chat.id, 
                message.message_id if (message.text is None or not message.text.startswith('/msg_id')) else int(message.text.split()[1]), 
                reply_markup=keyboard.support_keyboard_2()
            )

            await message.bot.set_message_reaction(message.chat.id, message.message_id, [{"type": "emoji", "emoji": "👍"}])

            file_id = None
            file_type = None

            if message.photo is not None:
                file_id = message.photo[-1].file_id
                file_type = 'photo'
            elif message.video is not None:
                file_id = message.video.file_id
                file_type = 'video'
            elif message.document:
                file_id = message.document.file_id
                file_type = 'document'
            elif message.audio:
                file_id = message.audio.file_id
                file_type = 'audio'
            elif message.voice:
                file_id = message.voice.file_id
                file_type = 'voice'
            elif message.video_note:
                file_id = message.video_note.file_id
                file_type = 'video_note'

            db.add_to_support_messages(message.from_user.id, message.chat.id, message.html_text, file_id, file_type, False, time.time(), f"https://t.me/c/{-(message.chat.id+1000000000000)}/{message.message_id}")

            # ✅ сервер рассылает событие (источник истины)
            message_payload = {
                "type": "message",
                "text": message.html_text,
                "sender_id": "0",
                "unix_time": int(time.time()),
            }

            # отправляем ВСЕМ подключённым (включая отправителя)
            for ws in config.ws_connections_support[str(user_id)]:
                await ws.send_json(message_payload)
        except Exception as e:
            print(e)
            pass

        return

    if message.chat.id not in chat_ids_list or message.reply_to_message is None:
        return
    
    if message.reply_to_message.text.__contains__('(Техническая информация:'):
        technical_info = message.reply_to_message.text.split("Техническая информация: ")[-1].split(")")[0].split("_")
        homework_data = db.get_homework_by_lesson_id(technical_info[0], technical_info[1])
        print(f"Использую новый алгоритм: {technical_info}")
    else:
        return
    
    if len(homework_data) == 0:
        return
    
    check_unix_time = 0 if homework_data[0]["check_time"] is None or len(homework_data[0]["check_time"]) == 0 else datetime.datetime.strptime(homework_data[0]["check_time"], "%Y-%m-%d %H:%M:%S").timestamp()
    print(check_unix_time)

    if time.time() - check_unix_time < 10: # На случай, если сообщение от трекера сплитанулось
        db.edit_homework(homework_data[0]["homework_id"], status=homework_data[0]["status"], comment=f'{homework_data[0]["comment"]}\n{message.text}', check_time=homework_data[0]["check_time"])
        print(f"Сообщение от трекера сплитанулось. ID: {homework_data[0]['homework_id']}")
        return

    if message.text.startswith('+') and len(message.text.splitlines()) == 1: # Трекер принял задание
        user_name = ''

        try:
            user_data = await message.bot.get_chat(int(homework_data[0]['user_data'].split()[-1]))
            user_name = user_data.first_name
        except:
            pass

        db.edit_homework(homework_data[0]["homework_id"], status='✅', comment=f'{user_name}, ты классно выполнила домашние задание. Корректировок нет, я принимаю его.', check_time=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        await send_congratulation_message(message, int(homework_data[0]['lesson_id']), int(homework_data[0]['user_data'].split()[-1]))
    elif message.text.splitlines()[0].startswith('+'): # Трекер принял задание
        message_text = message.text[1:].strip()
        db.edit_homework(homework_data[0]["homework_id"], status='✅', comment=message_text, check_time=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    elif message.text == '1': # Трекер отправил на проверку
        db.edit_homework(homework_data[0]["homework_id"], status='На проверке', check_time=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    else:
        db.edit_homework(homework_data[0]["homework_id"], status='❌', comment=message.text, check_time=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    try:
        await message.bot.set_message_reaction(message.chat.id, message.message_id, [{"type": "emoji", "emoji": "👍"}])
    except Exception as reaction_error:
        # Фолбэк для чатов, где реакции недоступны/запрещены
        try:
            await message.reply("✅ Принято")
        except Exception:
            pass

        try:
            await message.bot.send_message(
                config.LOG_CHAT_ID,
                f"Не удалось поставить реакцию в чате {message.chat.id} на сообщение {message.message_id}: {reaction_error}"
            )
        except Exception:
            pass

    if message.text.startswith('+'):
        await send_congratulation_message(message, int(homework_data[0]['lesson_id']), int(homework_data[0]['user_data'].split()[-1]))

    await message.bot.send_message(
        int(homework_data[0]['user_data'].split()[-1]),
        "Твой персональный трекер проверил твое задание🔥 Нажми на кнопку, чтобы посмотреть обратную связь👇🏻",
        reply_markup=keyboard.open_lesson_keyboard(homework_data[0]['lesson_id']) if message.text.startswith('+') else keyboard.get_last_solution_keyboard(int(homework_data[0]['lesson_id']))
    )

    # Удаление дублей ДЗ
    if len(homework_data) > 1:
        for i in range(1, len(homework_data)):
            try:
                db.delete_homework_by_homework_id(homework_data[i]["homework_id"])
                await message.bot.send_message(config.LOG_CHAT_ID, f'Удален дубль ДЗ: {homework_data[i]["homework_id"]}')
            except:
                pass
